"""
DOCX Resume Parser

Handles extraction and parsing of DOCX resumes using python-docx.
"""

import logging

logger = logging.getLogger(__name__)

# Optional dependency with graceful degradation
try:
    from docx import Document

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not available. DOCX parsing will be disabled.")


class DOCXParser:
    """Handles DOCX resume parsing and text extraction."""

    def can_parse(self, file_path: str) -> bool:
        """Check if this parser can handle the given file."""
        return HAS_DOCX and file_path.lower().endswith(".docx")

    def extract_text(self, file_path: str) -> str | None:
        """Extract text content from DOCX resume."""
        if not self.can_parse(file_path):
            return None

        try:
            doc = Document(file_path)
            text_content = []

            # Extract paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)

            return "\n".join(text_content) if text_content else None

        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            return None

    def extract_metadata(self, file_path: str) -> dict:
        """Extract metadata from DOCX resume."""
        if not self.can_parse(file_path):
            return {}

        try:
            doc = Document(file_path)
            core_props = doc.core_properties

            return {
                "author": core_props.author,
                "title": core_props.title,
                "subject": core_props.subject,
                "created": core_props.created,
                "modified": core_props.modified,
                "format": "docx",
            }
        except Exception as e:
            logger.error(f"Failed to extract metadata from DOCX {file_path}: {e}")
            return {}
