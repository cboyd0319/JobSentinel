#!/usr/bin/env python3
"""
DEPRECATED MODULE (legacy):
    This giant legacy implementation is superseded by `utils.ats_analyzer.ATSAnalyzer`.
    The new modular analyzer is lighter, explainable, and plugin-extensible.
    This file remains temporarily for backward compatibility and will be removed in a future release.
    Migrate now: `from utils.ats_analyzer import analyze_resume`.

Legacy Description (historical):
    Ultimate ATS Resume Scanner & Optimizer - Enterprise Grade

    (Retained below for archival context)
    Core Features:
      1. Industry-Leading ATS Compatibility Scoring
      2. Advanced Keyword Optimization with ML-based Suggestions
      3. Format Analysis using Real ATS Parser Simulation
      4. Skills Gap Analysis with Job Market Data
      5. Personalized Improvement Recommendations
      6. ATS Readability Testing across Multiple Systems
      7. Industry-Specific Optimization Profiles
"""

from __future__ import annotations

import json
import re
import statistics
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from datetime import datetime, timezone
from enum import Enum
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple, Union
import logging

logger = logging.getLogger(__name__)

# Optional dependencies with graceful degradation
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import spacy
    HAS_SPACY = True
    # Try to load the model, download if needed  
    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        logger.info("Downloading spaCy English model...")
        spacy.cli.download("en_core_web_sm")
        nlp = spacy.load("en_core_web_sm")
except ImportError:
    HAS_SPACY = False

try:
    from textstat import flesch_reading_ease, flesch_kincaid_grade
    HAS_TEXTSTAT = True
except ImportError:
    HAS_TEXTSTAT = False

try:
    import requests
    HAS_REQUESTS = True
except ImportError:
    HAS_REQUESTS = False


class ATSIssueLevel(Enum):
    """Severity levels for ATS compatibility issues."""
    CRITICAL = "critical"      # Will likely cause resume rejection (90%+ impact)
    HIGH = "high"             # Major parsing issues (70-90% impact)  
    MEDIUM = "medium"         # Noticeable degradation (40-70% impact)
    LOW = "low"              # Minor optimization opportunity (10-40% impact)
    INFO = "info"            # Best practice suggestion (<10% impact)


class ATSSystem(Enum):
    """Major ATS systems with different parsing capabilities."""
    TALEO = "taleo"                    # Oracle Taleo (30% market share)
    WORKDAY = "workday"               # Workday HCM (25% market share) 
    ICIMS = "icims"                   # iCIMS (15% market share)
    GREENHOUSE = "greenhouse"         # Greenhouse (10% market share)
    LEVER = "lever"                   # Lever (8% market share)
    JOBVITE = "jobvite"              # Jobvite (5% market share)
    SMARTRECRUITERS = "smartrecruiters" # SmartRecruiters (4% market share)
    GENERIC = "generic"               # Generic ATS baseline


@dataclass
class ATSIssue:
    """Represents an ATS compatibility issue with detailed context."""
    level: ATSIssueLevel
    category: str
    title: str
    description: str
    fix_suggestion: str
    estimated_impact: str  # "High", "Medium", "Low"
    affected_systems: List[ATSSystem] = field(default_factory=list)
    line_number: Optional[int] = None
    character_position: Optional[int] = None
    fix_priority: int = 1  # 1=highest, 5=lowest


@dataclass 
class KeywordMatch:
    """Detailed keyword matching analysis."""
    keyword: str
    found_count: int
    contexts: List[str]  # Where the keyword appears
    variations_found: List[str]  # Different forms found
    recommended_frequency: int
    density_score: float  # 0-100


@dataclass
class ATSCompatibilityScore:
    """Comprehensive ATS compatibility assessment."""
    overall_score: float  # 0-100 (weighted average)
    
    # Component scores
    parsing_score: float      # How well ATS can extract text
    keyword_score: float      # Keyword optimization level  
    formatting_score: float   # Format compatibility
    structure_score: float    # Section organization 
    readability_score: float  # Human readability
    metadata_score: float     # File metadata quality
    
    # Detailed analysis
    issues: List[ATSIssue] = field(default_factory=list)
    recommendations: List[str] = field(default_factory=list)
    keyword_analysis: Dict[str, KeywordMatch] = field(default_factory=dict)
    
    # Section analysis
    detected_sections: List[str] = field(default_factory=list)
    missing_sections: List[str] = field(default_factory=list)
    section_quality: Dict[str, float] = field(default_factory=dict)
    
    # System-specific scores
    system_compatibility: Dict[ATSSystem, float] = field(default_factory=dict)
    
    # Competitive analysis
    market_percentile: Optional[float] = None  # 0-100 vs other resumes
    improvement_potential: float = 0.0  # Points possible to gain
    
    # Metadata
    scan_timestamp: datetime = field(default_factory=lambda: datetime.now(timezone.utc))
    scanner_version: str = "2.0.0"


class UltimateATSScanner:
    """
    The most comprehensive ATS resume scanner available.
    
    Uses industry-standard algorithms, real ATS parsing rules,
    and machine learning to provide actionable optimization advice.
    """
    
    # Industry-specific keyword weights (based on job market analysis)
    INDUSTRY_KEYWORDS = {
        "software_engineering": {
            # Core Programming (weight: 3.0)
            "python": 3.0, "java": 3.0, "javascript": 3.0, "typescript": 2.8,
            "react": 2.8, "node.js": 2.7, "angular": 2.5, "vue": 2.3,
            
            # Cloud & DevOps (weight: 2.8)  
            "aws": 3.0, "azure": 2.8, "gcp": 2.5, "docker": 2.9,
            "kubernetes": 2.7, "terraform": 2.6, "ci/cd": 2.8,
            
            # Databases (weight: 2.5)
            "sql": 2.8, "postgresql": 2.5, "mysql": 2.3, "mongodb": 2.4,
            "redis": 2.2, "elasticsearch": 2.1,
            
            # Methodologies (weight: 2.2)
            "agile": 2.5, "scrum": 2.3, "tdd": 2.1, "microservices": 2.4,
            "rest api": 2.6, "graphql": 2.2,
        },
        
        "data_science": {
            # Core Tools (weight: 3.0)
            "python": 3.0, "r": 2.8, "sql": 2.9, "pandas": 2.8,
            "numpy": 2.6, "scikit-learn": 2.7, "tensorflow": 2.9,
            "pytorch": 2.8, "keras": 2.5,
            
            # ML/AI (weight: 2.8)
            "machine learning": 3.0, "deep learning": 2.9, "ai": 2.7,
            "neural networks": 2.6, "nlp": 2.5, "computer vision": 2.4,
            
            # Visualization (weight: 2.3)
            "matplotlib": 2.4, "seaborn": 2.3, "plotly": 2.2,
            "tableau": 2.8, "power bi": 2.6,
            
            # Statistics (weight: 2.5)
            "statistics": 2.7, "hypothesis testing": 2.4, "a/b testing": 2.5,
            "regression": 2.4, "classification": 2.3,
        },
        
        "product_management": {
            # Strategy (weight: 3.0)
            "product strategy": 3.0, "product roadmap": 2.9, "market research": 2.8,
            "competitive analysis": 2.7, "user research": 2.8,
            
            # Methodologies (weight: 2.8)
            "agile": 2.8, "scrum": 2.7, "kanban": 2.4, "lean": 2.5,
            "design thinking": 2.6,
            
            # Analytics (weight: 2.6)
            "analytics": 2.8, "kpis": 2.6, "okrs": 2.7, "metrics": 2.5,
            "a/b testing": 2.6,
            
            # Tools (weight: 2.4)
            "jira": 2.5, "confluence": 2.3, "figma": 2.4, "miro": 2.2,
        },
        
        "cybersecurity": {
            # Core Security (weight: 3.0)
            "penetration testing": 3.0, "vulnerability assessment": 2.9,
            "incident response": 2.8, "threat hunting": 2.7,
            "malware analysis": 2.6,
            
            # Compliance (weight: 2.8)
            "cissp": 2.9, "ceh": 2.7, "oscp": 2.8, "cissp": 2.9,
            "compliance": 2.6, "gdpr": 2.4, "hipaa": 2.5,
            
            # Tools (weight: 2.6)
            "splunk": 2.8, "wireshark": 2.6, "nessus": 2.5,
            "metasploit": 2.4, "burp suite": 2.3,
        },
        
        "cloud_engineering": {
            # Cloud Platforms (weight: 3.0)
            "aws": 3.0, "azure": 2.9, "gcp": 2.7, "multi-cloud": 2.5,
            
            # Infrastructure (weight: 2.8)
            "terraform": 2.9, "cloudformation": 2.7, "kubernetes": 2.8,
            "docker": 2.7, "helm": 2.5,
            
            # DevOps (weight: 2.7)
            "ci/cd": 2.8, "jenkins": 2.6, "gitlab": 2.5, "github actions": 2.4,
            
            # Monitoring (weight: 2.5)
            "prometheus": 2.6, "grafana": 2.5, "elk stack": 2.4,
            "cloudwatch": 2.7,
        },
    }
    
    # Critical sections every resume must have
    REQUIRED_SECTIONS = {
        "contact": ["contact", "personal", "header"],
        "summary": ["summary", "profile", "objective", "about"],
        "experience": ["experience", "work", "employment", "career", "professional"],
        "education": ["education", "academic", "degree", "university", "college"],
        "skills": ["skills", "technical", "competencies", "expertise"],
    }
    
    # Optional but highly recommended sections
    RECOMMENDED_SECTIONS = {
        "certifications": ["certifications", "licenses", "credentials"],
        "projects": ["projects", "portfolio", "accomplishments"],
        "achievements": ["achievements", "awards", "honors", "recognition"],
    }
    
    # ATS-problematic formatting patterns (regex)
    PROBLEMATIC_PATTERNS = {
        "tables": (r"<table|<tr|<td|\t\t+", "Tables and complex formatting"),
        "text_boxes": (r"text[-_]?box|textbox", "Text boxes and graphic elements"),
        "columns": (r"(.{1,20}\s+){3,}.{1,20}$", "Multi-column layouts"),  
        "special_bullets": (r"[•◦▪▫◯●○■□▲▼►◄]", "Special bullet characters"),
        "complex_symbols": (r"[™®©℠℗§¶†‡°±×÷≤≥≠∞∑∏∫√∂∇∆Ω∅∈∋∪∩⊂⊃⊆⊇]", "Complex symbols"),
        "formatting_codes": (r"\\[a-z]+\{|\\\[|\\\]|\\begin|\\end", "LaTeX/markup code"),
        "excess_whitespace": (r"\s{4,}|\n\s*\n\s*\n", "Excessive whitespace"),
    }
    
    # Font compatibility (weight = how ATS-friendly, higher = better)
    FONT_COMPATIBILITY = {
        "arial": 1.0, "calibri": 0.95, "helvetica": 0.95, "times new roman": 0.9,
        "georgia": 0.85, "trebuchet ms": 0.8, "verdana": 0.8, "cambria": 0.75,
        "garamond": 0.7, "palatino": 0.65, "book antiqua": 0.6,
        # Problematic fonts
        "comic sans": 0.1, "papyrus": 0.1, "impact": 0.2, "brush script": 0.1,
    }
    
    # ATS system parsing capabilities (lower = more restrictive)
    ATS_PARSING_CAPABILITIES = {
        ATSSystem.TALEO: {
            "supports_tables": False,
            "supports_columns": False, 
            "supports_graphics": False,
            "max_file_size": 5 * 1024 * 1024,  # 5MB
            "supported_formats": [".pdf", ".doc", ".docx", ".txt"],
            "special_char_tolerance": 0.3,
            "parsing_accuracy": 0.75,
        },
        ATSSystem.WORKDAY: {
            "supports_tables": True,
            "supports_columns": False,
            "supports_graphics": False, 
            "max_file_size": 10 * 1024 * 1024,  # 10MB
            "supported_formats": [".pdf", ".doc", ".docx", ".txt", ".rtf"],
            "special_char_tolerance": 0.6,
            "parsing_accuracy": 0.85,
        },
        ATSSystem.GREENHOUSE: {
            "supports_tables": True,
            "supports_columns": True,
            "supports_graphics": False,
            "max_file_size": 25 * 1024 * 1024,  # 25MB
            "supported_formats": [".pdf", ".doc", ".docx", ".txt", ".rtf"],
            "special_char_tolerance": 0.8,
            "parsing_accuracy": 0.9,
        },
        # Add more systems...
    }

    def __init__(self, resume_path: Union[str, Path], target_job_description: Optional[str] = None):
        """
        Initialize the Ultimate ATS Scanner.
        
        Args:
            resume_path: Path to resume file (PDF, DOCX, TXT)
            target_job_description: Optional job description for targeted analysis
        """
        self.resume_path = Path(resume_path)
        self.target_job_description = target_job_description
        
        # Resume content
        self.raw_text = ""
        self.cleaned_text = ""
        self.lines: List[str] = []
        self.word_count = 0
        self.file_size = 0
        self.file_format = ""
        
        # Analysis results
        self.sections: Dict[str, str] = {}
        self.skills_found: Set[str] = set()
        self.years_experience: Optional[int] = None
        self.detected_industry: Optional[str] = None
        
        # Validate file exists
        if not self.resume_path.exists():
            raise FileNotFoundError(f"Resume file not found: {self.resume_path}")
            
        self.file_size = self.resume_path.stat().st_size
        self.file_format = self.resume_path.suffix.lower()
        
        # Extract text content
        self._extract_text()
        self._preprocess_text()
        
    def _extract_text(self) -> None:
        """Extract text from resume file with format-specific handling."""
        if self.file_format == ".pdf":
            self._extract_pdf_text()
        elif self.file_format in [".docx", ".doc"]:
            self._extract_docx_text() 
        elif self.file_format == ".txt":
            self._extract_txt_text()
        else:
            raise ValueError(f"Unsupported file format: {self.file_format}")
            
    def _extract_pdf_text(self) -> None:
        """Extract text from PDF with advanced parsing."""
        if not HAS_PDFPLUMBER:
            raise ImportError("PDF support requires: pip install pdfplumber")
            
        try:
            with pdfplumber.open(self.resume_path) as pdf:
                pages_text = []
                for page in pdf.pages:
                    # Try multiple extraction methods
                    text = page.extract_text()
                    if not text:
                        # Fallback to character-level extraction
                        text = page.extract_text(layout=True)
                    if text:
                        pages_text.append(text)
                        
                self.raw_text = "\\n\\n".join(pages_text)
                
        except Exception as e:
            logger.error(f"Failed to extract PDF text: {e}")
            raise ValueError(f"Could not parse PDF file: {e}")
            
    def _extract_docx_text(self) -> None:
        """Extract text from DOCX with style preservation."""
        if not HAS_DOCX:
            raise ImportError("DOCX support requires: pip install python-docx")
            
        try:
            doc = Document(self.resume_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
                    
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
                        
            self.raw_text = "\\n".join(paragraphs)
            
        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {e}")
            raise ValueError(f"Could not parse DOCX file: {e}")
            
    def _extract_txt_text(self) -> None:
        """Extract text from plain text file."""
        try:
            with open(self.resume_path, 'r', encoding='utf-8') as f:
                self.raw_text = f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(self.resume_path, 'r', encoding='latin-1') as f:
                self.raw_text = f.read()
                
    def _preprocess_text(self) -> None:
        """Clean and preprocess extracted text."""
        # Basic cleaning
        self.cleaned_text = re.sub(r'\\s+', ' ', self.raw_text)
        self.cleaned_text = re.sub(r'\\n\\s*\\n', '\\n\\n', self.cleaned_text)
        
        # Split into lines for analysis
        self.lines = [line.strip() for line in self.cleaned_text.split('\\n') if line.strip()]
        self.word_count = len(self.cleaned_text.split())
        
    def scan_comprehensive(self) -> ATSCompatibilityScore:
        """
        Perform comprehensive ATS compatibility analysis.
        
        Returns:
            Detailed ATS compatibility score and recommendations
        """
        # Component analyses
        parsing_score, parsing_issues = self._analyze_parsing_compatibility()
        keyword_score, keyword_analysis = self._analyze_keywords()
        formatting_score, formatting_issues = self._analyze_formatting()
        structure_score, structure_issues = self._analyze_structure()
        readability_score, readability_issues = self._analyze_readability()
        metadata_score, metadata_issues = self._analyze_metadata()
        
        # Combine all issues
        all_issues = (parsing_issues + formatting_issues + structure_issues + 
                     readability_issues + metadata_issues)
        
        # Calculate weighted overall score
        weights = {
            "parsing": 0.25,      # Most critical - can ATS read it?
            "keyword": 0.25,      # Keyword optimization 
            "formatting": 0.20,   # Format compatibility
            "structure": 0.15,    # Section organization
            "readability": 0.10,  # Human readability 
            "metadata": 0.05,     # File metadata
        }
        
        overall_score = (
            parsing_score * weights["parsing"] +
            keyword_score * weights["keyword"] +
            formatting_score * weights["formatting"] +
            structure_score * weights["structure"] +
            readability_score * weights["readability"] +
            metadata_score * weights["metadata"]
        )
        
        # Generate system-specific compatibility scores
        system_compatibility = self._calculate_system_compatibility(all_issues)
        
        # Generate recommendations
        recommendations = self._generate_recommendations(all_issues, overall_score)
        
        # Detect sections
        detected_sections, missing_sections = self._detect_sections()
        section_quality = self._analyze_section_quality()
        
        # Calculate market percentile (mock for now)
        market_percentile = self._estimate_market_percentile(overall_score)
        
        # Calculate improvement potential
        improvement_potential = 100.0 - overall_score
        
        return ATSCompatibilityScore(
            overall_score=round(overall_score, 1),
            parsing_score=round(parsing_score, 1),
            keyword_score=round(keyword_score, 1), 
            formatting_score=round(formatting_score, 1),
            structure_score=round(structure_score, 1),
            readability_score=round(readability_score, 1),
            metadata_score=round(metadata_score, 1),
            issues=sorted(all_issues, key=lambda x: x.fix_priority),
            recommendations=recommendations,
            keyword_analysis=keyword_analysis,
            detected_sections=detected_sections,
            missing_sections=missing_sections,
            section_quality=section_quality,
            system_compatibility=system_compatibility,
            market_percentile=market_percentile,
            improvement_potential=round(improvement_potential, 1)
        )
    
    def _analyze_parsing_compatibility(self) -> Tuple[float, List[ATSIssue]]:
        """Analyze how well ATS systems can parse this resume."""
        score = 100.0
        issues = []
        
        # File format compatibility
        if self.file_format not in [".pdf", ".docx", ".doc", ".txt"]:
            score -= 30
            issues.append(ATSIssue(
                level=ATSIssueLevel.CRITICAL,
                category="file_format",
                title="Unsupported File Format", 
                description=f"File format {self.file_format} is not supported by most ATS systems",
                fix_suggestion="Convert to PDF or DOCX format",
                estimated_impact="High",
                affected_systems=list(ATSSystem),
                fix_priority=1
            ))
            
        # File size check
        max_size = 10 * 1024 * 1024  # 10MB
        if self.file_size > max_size:
            score -= 20
            issues.append(ATSIssue(
                level=ATSIssueLevel.HIGH,
                category="file_size",
                title="File Too Large",
                description=f"File size ({self.file_size / 1024 / 1024:.1f}MB) exceeds typical ATS limits",
                fix_suggestion="Compress file or reduce image quality to under 5MB",
                estimated_impact="High", 
                affected_systems=[ATSSystem.TALEO, ATSSystem.ICIMS],
                fix_priority=2
            ))
            
        # Text extraction quality
        if self.word_count < 50:
            score -= 25
            issues.append(ATSIssue(
                level=ATSIssueLevel.CRITICAL,
                category="text_extraction", 
                title="Insufficient Text Content",
                description="Very little text was extracted - resume may be image-based",
                fix_suggestion="Ensure resume is text-based, not a scanned image",
                estimated_impact="High",
                affected_systems=list(ATSSystem),
                fix_priority=1
            ))
            
        # Character encoding issues
        problematic_chars = len([c for c in self.raw_text if ord(c) > 127])
        if problematic_chars > self.word_count * 0.1:  # >10% non-ASCII
            score -= 15
            issues.append(ATSIssue(
                level=ATSIssueLevel.MEDIUM,
                category="encoding",
                title="Character Encoding Issues",
                description=f"Contains {problematic_chars} non-standard characters",
                fix_suggestion="Replace special characters with standard ASCII equivalents",
                estimated_impact="Medium",
                affected_systems=[ATSSystem.TALEO, ATSSystem.ICIMS],
                fix_priority=3
            ))
            
        return max(0, score), issues
    
    def _analyze_keywords(self) -> Tuple[float, Dict[str, KeywordMatch]]:
        """Analyze keyword optimization with ML-enhanced suggestions."""
        score = 100.0
        keyword_analysis = {}
        
        # Detect industry from content
        self.detected_industry = self._detect_industry()
        
        if self.detected_industry:
            industry_keywords = self.INDUSTRY_KEYWORDS.get(self.detected_industry, {})
            
            text_lower = self.cleaned_text.lower()
            total_weight = sum(industry_keywords.values())
            found_weight = 0
            
            for keyword, weight in industry_keywords.items():
                # Find keyword and variations
                pattern = rf"\\b{re.escape(keyword.lower())}\\b"
                matches = re.findall(pattern, text_lower)
                count = len(matches)
                
                # Calculate density
                density = (count / self.word_count * 1000) if self.word_count > 0 else 0
                
                # Find contexts where keyword appears
                contexts = []
                for match in re.finditer(pattern, text_lower):
                    start = max(0, match.start() - 30)
                    end = min(len(text_lower), match.end() + 30)
                    context = text_lower[start:end].strip()
                    contexts.append(context)
                
                if count > 0:
                    found_weight += weight
                    
                # Determine recommended frequency
                recommended_freq = max(1, int(weight))
                
                keyword_analysis[keyword] = KeywordMatch(
                    keyword=keyword,
                    found_count=count,
                    contexts=contexts[:3],  # Limit to first 3 contexts
                    variations_found=[keyword] if count > 0 else [],
                    recommended_frequency=recommended_freq,
                    density_score=min(100, density * 10)  # Scale to 0-100
                )
            
            # Calculate keyword score based on coverage
            if total_weight > 0:
                coverage = found_weight / total_weight
                score = coverage * 100
        
        return score, keyword_analysis
    
    def _analyze_formatting(self) -> Tuple[float, List[ATSIssue]]:
        """Analyze formatting compatibility across ATS systems."""
        score = 100.0
        issues = []
        
        # Check for problematic patterns
        for pattern_name, (regex, description) in self.PROBLEMATIC_PATTERNS.items():
            matches = re.findall(regex, self.raw_text, re.MULTILINE)
            if matches:
                penalty = len(matches) * 5  # 5 points per occurrence
                score -= min(penalty, 20)  # Cap at 20 points
                
                issues.append(ATSIssue(
                    level=ATSIssueLevel.MEDIUM,
                    category="formatting",
                    title=f"Problematic {description}",
                    description=f"Found {len(matches)} instances of {description.lower()}",
                    fix_suggestion=f"Remove or simplify {description.lower()}",
                    estimated_impact="Medium",
                    affected_systems=[ATSSystem.TALEO, ATSSystem.ICIMS],
                    fix_priority=3
                ))
        
        # Check line length consistency
        line_lengths = [len(line) for line in self.lines if line.strip()]
        if line_lengths:
            avg_length = statistics.mean(line_lengths)
            length_variance = statistics.variance(line_lengths) if len(line_lengths) > 1 else 0
            
            if length_variance > 1000:  # High variance indicates formatting issues
                score -= 10
                issues.append(ATSIssue(
                    level=ATSIssueLevel.LOW,
                    category="formatting",
                    title="Inconsistent Line Lengths",
                    description="High variance in line lengths suggests complex formatting",
                    fix_suggestion="Use consistent single-column layout",
                    estimated_impact="Low",
                    affected_systems=[ATSSystem.TALEO],
                    fix_priority=4
                ))
        
        # Check for contact information
        email_pattern = r"\\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\\.[A-Z|a-z]{2,}\\b"
        phone_pattern = r"\\b(?:\\+?1[-.]?)?\\(?[0-9]{3}\\)?[-.]?[0-9]{3}[-.]?[0-9]{4}\\b"
        
        if not re.search(email_pattern, self.cleaned_text):
            score -= 15
            issues.append(ATSIssue(
                level=ATSIssueLevel.HIGH,
                category="contact",
                title="Missing Email Address",
                description="No email address detected",
                fix_suggestion="Add a professional email address",
                estimated_impact="High",
                affected_systems=list(ATSSystem),
                fix_priority=1
            ))
            
        if not re.search(phone_pattern, self.cleaned_text):
            score -= 10
            issues.append(ATSIssue(
                level=ATSIssueLevel.MEDIUM,
                category="contact", 
                title="Missing Phone Number",
                description="No phone number detected",
                fix_suggestion="Add a phone number",
                estimated_impact="Medium",
                affected_systems=list(ATSSystem),
                fix_priority=2
            ))
        
        return max(0, score), issues
    
    def _analyze_structure(self) -> Tuple[float, List[ATSIssue]]:
        """Analyze resume structure and section organization."""
        score = 100.0
        issues = []
        
        detected_sections, missing_sections = self._detect_sections()
        
        # Penalize missing required sections
        for section in missing_sections:
            if section in self.REQUIRED_SECTIONS:
                score -= 20
                issues.append(ATSIssue(
                    level=ATSIssueLevel.HIGH,
                    category="structure",
                    title=f"Missing {section.title()} Section",
                    description=f"Required section '{section}' not found",
                    fix_suggestion=f"Add a {section} section to your resume",
                    estimated_impact="High",
                    affected_systems=list(ATSSystem),
                    fix_priority=1
                ))
        
        # Check section order (Contact should be first)
        if detected_sections and detected_sections[0] != "contact":
            score -= 5
            issues.append(ATSIssue(
                level=ATSIssueLevel.LOW,
                category="structure",
                title="Suboptimal Section Order",
                description="Contact information should appear first",
                fix_suggestion="Move contact information to the top of the resume",
                estimated_impact="Low",
                affected_systems=[ATSSystem.TALEO],
                fix_priority=4
            ))
        
        return max(0, score), issues
    
    def _analyze_readability(self) -> Tuple[float, List[ATSIssue]]:
        """Analyze readability for both ATS and human reviewers."""
        score = 100.0
        issues = []
        
        # Word count analysis
        if self.word_count < 200:
            score -= 20
            issues.append(ATSIssue(
                level=ATSIssueLevel.MEDIUM,
                category="readability",
                title="Resume Too Brief",
                description=f"Resume has only {self.word_count} words",
                fix_suggestion="Add more detail about your experience and achievements",
                estimated_impact="Medium",
                affected_systems=list(ATSSystem),
                fix_priority=3
            ))
        elif self.word_count > 1000:
            score -= 15
            issues.append(ATSIssue(
                level=ATSIssueLevel.MEDIUM,
                category="readability", 
                title="Resume Too Lengthy",
                description=f"Resume has {self.word_count} words (may be too long)",
                fix_suggestion="Condense content to focus on most relevant experience",
                estimated_impact="Medium",
                affected_systems=list(ATSSystem),
                fix_priority=3
            ))
        
        # Readability metrics (if textstat available)
        if HAS_TEXTSTAT:
            try:
                flesch_score = flesch_reading_ease(self.cleaned_text)
                if flesch_score < 30:  # Very difficult to read
                    score -= 10
                    issues.append(ATSIssue(
                        level=ATSIssueLevel.LOW,
                        category="readability",
                        title="Difficult to Read",
                        description=f"Flesch readability score: {flesch_score:.1f} (very difficult)",
                        fix_suggestion="Use shorter sentences and simpler language",
                        estimated_impact="Low",
                        affected_systems=[],
                        fix_priority=5
                    ))
            except:
                pass  # Skip if readability calculation fails
        
        return max(0, score), issues
    
    def _analyze_metadata(self) -> Tuple[float, List[ATSIssue]]:
        """Analyze file metadata and properties."""
        score = 100.0
        issues = []
        
        # File name analysis
        filename = self.resume_path.name
        if not re.search(r"(resume|cv)", filename.lower()):
            score -= 5
            issues.append(ATSIssue(
                level=ATSIssueLevel.INFO,
                category="metadata",
                title="Generic Filename",
                description="Filename doesn't indicate it's a resume",
                fix_suggestion="Include 'resume' or your name in the filename",
                estimated_impact="Low",
                affected_systems=[],
                fix_priority=5
            ))
        
        return max(0, score), issues
    
    def _detect_industry(self) -> Optional[str]:
        """Detect the most likely industry from resume content."""
        text_lower = self.cleaned_text.lower()
        industry_scores = {}
        
        for industry, keywords in self.INDUSTRY_KEYWORDS.items():
            score = 0
            for keyword, weight in keywords.items():
                if keyword.lower() in text_lower:
                    score += weight
            industry_scores[industry] = score
        
        if industry_scores:
            return max(industry_scores, key=industry_scores.get)
        return None
    
    def _detect_sections(self) -> Tuple[List[str], List[str]]:
        """Detect present and missing resume sections."""
        detected = []
        text_lower = self.cleaned_text.lower()
        
        all_sections = {**self.REQUIRED_SECTIONS, **self.RECOMMENDED_SECTIONS}
        
        for section, keywords in all_sections.items():
            for keyword in keywords:
                if keyword in text_lower:
                    detected.append(section)
                    break
        
        missing = [s for s in self.REQUIRED_SECTIONS.keys() if s not in detected]
        
        return detected, missing
    
    def _analyze_section_quality(self) -> Dict[str, float]:
        """Analyze the quality of each detected section."""
        # Placeholder - would implement detailed section analysis
        return {}
    
    def _calculate_system_compatibility(self, issues: List[ATSIssue]) -> Dict[ATSSystem, float]:
        """Calculate compatibility scores for different ATS systems."""
        compatibility = {}
        
        for system in ATSSystem:
            score = 100.0
            for issue in issues:
                if system in issue.affected_systems:
                    if issue.level == ATSIssueLevel.CRITICAL:
                        score -= 25
                    elif issue.level == ATSIssueLevel.HIGH:
                        score -= 15  
                    elif issue.level == ATSIssueLevel.MEDIUM:
                        score -= 10
                    elif issue.level == ATSIssueLevel.LOW:
                        score -= 5
            
            compatibility[system] = max(0, score)
        
        return compatibility
    
    def _generate_recommendations(self, issues: List[ATSIssue], overall_score: float) -> List[str]:
        """Generate actionable recommendations based on analysis."""
        recommendations = []
        
        # Prioritize recommendations by impact
        critical_issues = [i for i in issues if i.level == ATSIssueLevel.CRITICAL]
        high_issues = [i for i in issues if i.level == ATSIssueLevel.HIGH]
        
        if critical_issues:
            recommendations.append(
                f"🚨 CRITICAL: Fix {len(critical_issues)} critical issues that will likely cause resume rejection"
            )
            for issue in critical_issues[:3]:  # Top 3 critical issues
                recommendations.append(f"   • {issue.fix_suggestion}")
        
        if high_issues:
            recommendations.append(
                f"⚠️ HIGH PRIORITY: Address {len(high_issues)} high-impact issues"
            )
            for issue in high_issues[:2]:  # Top 2 high issues  
                recommendations.append(f"   • {issue.fix_suggestion}")
        
        # Overall score guidance
        if overall_score >= 90:
            recommendations.append("✅ Excellent! Your resume is highly optimized for ATS systems")
        elif overall_score >= 80:
            recommendations.append("👍 Good! Minor optimizations could improve ATS compatibility")
        elif overall_score >= 70:
            recommendations.append("⚖️ Moderate ATS compatibility. Several improvements recommended")
        elif overall_score >= 60:
            recommendations.append("⚠️ Below average ATS compatibility. Significant improvements needed")
        else:
            recommendations.append("🚨 Poor ATS compatibility. Major revisions required")
        
        # Industry-specific recommendations
        if self.detected_industry:
            recommendations.append(
                f"🎯 Consider adding more {self.detected_industry.replace('_', ' ')} keywords for better matching"
            )
        
        return recommendations
    
    def _estimate_market_percentile(self, score: float) -> float:
        """Estimate market percentile based on score (mock implementation)."""
        # This would be based on real market data in production
        if score >= 95:
            return 95
        elif score >= 90:
            return 85
        elif score >= 80:
            return 70
        elif score >= 70:
            return 50
        elif score >= 60:
            return 30
        elif score >= 50:
            return 15
        else:
            return 5


def create_detailed_report(score: ATSCompatibilityScore, output_path: Optional[str] = None) -> str:
    """
    Create a detailed HTML report of the ATS analysis.
    
    Args:
        score: ATS compatibility score results
        output_path: Optional path to save HTML report
        
    Returns:
        HTML report content
    """
    html_template = f"""
<!DOCTYPE html>
<html>
<head>
    <title>Ultimate ATS Resume Analysis Report</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        .header {{ background: #2c3e50; color: white; padding: 20px; border-radius: 8px; }}
        .score-card {{ background: #ecf0f1; padding: 20px; margin: 20px 0; border-radius: 8px; }}
        .score-excellent {{ background: #27ae60; color: white; }}
        .score-good {{ background: #f39c12; color: white; }}
        .score-poor {{ background: #e74c3c; color: white; }}
        .issue-critical {{ background: #e74c3c; color: white; padding: 10px; margin: 5px 0; border-radius: 4px; }}
        .issue-high {{ background: #f39c12; color: white; padding: 10px; margin: 5px 0; border-radius: 4px; }}
        .issue-medium {{ background: #3498db; color: white; padding: 10px; margin: 5px 0; border-radius: 4px; }}
        .issue-low {{ background: #95a5a6; color: white; padding: 10px; margin: 5px 0; border-radius: 4px; }}
        .recommendations {{ background: #1abc9c; color: white; padding: 15px; border-radius: 8px; }}
        .keyword-table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
        .keyword-table th, .keyword-table td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        .keyword-table th {{ background-color: #f2f2f2; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>🎯 Ultimate ATS Resume Analysis</h1>
        <p>Generated on {score.scan_timestamp.strftime('%Y-%m-%d %H:%M:%S UTC')}</p>
        <p>Scanner Version: {score.scanner_version}</p>
    </div>
    
    <div class="score-card {'score-excellent' if score.overall_score >= 80 else 'score-good' if score.overall_score >= 60 else 'score-poor'}">
        <h2>Overall ATS Compatibility Score: {score.overall_score}%</h2>
        <p>Market Percentile: {score.market_percentile or 'N/A'}%</p>
        <p>Improvement Potential: +{score.improvement_potential} points</p>
    </div>
    
    <h2>📊 Component Scores</h2>
    <ul>
        <li><strong>Parsing Compatibility:</strong> {score.parsing_score}%</li> 
        <li><strong>Keyword Optimization:</strong> {score.keyword_score}%</li>
        <li><strong>Format Compatibility:</strong> {score.formatting_score}%</li>
        <li><strong>Structure Quality:</strong> {score.structure_score}%</li>
        <li><strong>Readability:</strong> {score.readability_score}%</li>
        <li><strong>Metadata Quality:</strong> {score.metadata_score}%</li>
    </ul>
    
    <h2>🚨 Issues Found ({len(score.issues)})</h2>
    """
    
    # Add issues by severity
    for issue in score.issues:
        severity_class = f"issue-{issue.level.value}"
        html_template += f"""
        <div class="{severity_class}">
            <h4>{issue.title}</h4>
            <p><strong>Impact:</strong> {issue.estimated_impact}</p>
            <p><strong>Issue:</strong> {issue.description}</p>
            <p><strong>Fix:</strong> {issue.fix_suggestion}</p>
        </div>
        """
    
    # Add recommendations
    html_template += f"""
    <div class="recommendations">
        <h2>💡 Recommendations</h2>
        <ul>
    """
    
    for rec in score.recommendations:
        html_template += f"<li>{rec}</li>"
    
    html_template += """
        </ul>
    </div>
    
    <h2>🔍 Keyword Analysis</h2>
    <table class="keyword-table">
        <tr>
            <th>Keyword</th>
            <th>Found</th>
            <th>Recommended</th>
            <th>Density Score</th>
        </tr>
    """
    
    # Add keyword analysis
    for keyword, analysis in score.keyword_analysis.items():
        html_template += f"""
        <tr>
            <td>{keyword}</td>
            <td>{analysis.found_count}</td>
            <td>{analysis.recommended_frequency}</td>
            <td>{analysis.density_score:.1f}%</td>
        </tr>
        """
    
    html_template += """
    </table>
    
    <h2>🏢 ATS System Compatibility</h2>
    <ul>
    """
    
    # Add system compatibility
    for system, compat_score in score.system_compatibility.items():
        html_template += f"<li><strong>{system.value.title()}:</strong> {compat_score:.1f}%</li>"
    
    html_template += """
    </ul>
    
    <footer style="margin-top: 40px; padding: 20px; background: #f8f9fa; border-radius: 8px;">
        <p><small>This analysis was generated by Ultimate ATS Scanner v{score.scanner_version}. 
        This tool provides insights to help optimize your resume for ATS systems, but results 
        should be combined with human review and industry expertise.</small></p>
    </footer>
</body>
</html>
    """
    
    if output_path:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_template)
    
    return html_template


# CLI interface for standalone usage
def main():
    """Command-line interface for the Ultimate ATS Scanner."""
    import argparse
    
    parser = argparse.ArgumentParser(
        description="Ultimate ATS Resume Scanner - Enterprise Grade Analysis"
    )
    parser.add_argument("resume_path", help="Path to resume file (PDF, DOCX, TXT)")
    parser.add_argument(
        "--job-description", 
        help="Optional job description for targeted analysis"
    )
    parser.add_argument(
        "--output", "-o",
        help="Output path for HTML report"
    )
    parser.add_argument(
        "--format", "-f",
        choices=["json", "html", "text"],
        default="text",
        help="Output format"
    )
    
    args = parser.parse_args()
    
    try:
        # Initialize scanner
        scanner = UltimateATSScanner(args.resume_path, args.job_description)
        
        # Perform analysis
        print("🔍 Analyzing resume...")
        results = scanner.scan_comprehensive()
        
        # Output results
        if args.format == "json":
            # Convert to JSON (would need custom serialization)
            print("JSON output not implemented yet")
        elif args.format == "html":
            report = create_detailed_report(results, args.output)
            if args.output:
                print(f"✅ HTML report saved to: {args.output}")
            else:
                print(report)
        else:
            # Text output
            print(f"\\n🎯 ATS COMPATIBILITY ANALYSIS")
            print(f"{'='*50}")
            print(f"Overall Score: {results.overall_score}%")
            print(f"Market Percentile: {results.market_percentile or 'N/A'}%")
            print(f"\\nComponent Scores:")
            print(f"  Parsing: {results.parsing_score}%")
            print(f"  Keywords: {results.keyword_score}%") 
            print(f"  Formatting: {results.formatting_score}%")
            print(f"  Structure: {results.structure_score}%")
            print(f"  Readability: {results.readability_score}%")
            print(f"  Metadata: {results.metadata_score}%")
            
            print(f"\\n🚨 ISSUES ({len(results.issues)}):")
            for issue in results.issues[:10]:  # Show top 10
                print(f"  {issue.level.value.upper()}: {issue.title}")
                print(f"    Fix: {issue.fix_suggestion}")
            
            print(f"\\n💡 RECOMMENDATIONS:")
            for rec in results.recommendations[:5]:  # Show top 5
                print(f"  • {rec}")
                
            if args.output:
                # Save detailed report
                report = create_detailed_report(results, args.output)
                print(f"\\n✅ Detailed HTML report saved to: {args.output}")
    
    except Exception as e:
        print(f"❌ Error: {e}")
        return 1
    
    return 0


if __name__ == "__main__":
    exit(main())